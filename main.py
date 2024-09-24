import argparse
import datetime
import xml.etree.ElementTree as ET

from docx import Document
from docx.shared import Pt
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# Mapping note names and notes. 
MAP_NOTE = {
    'R': '0', 
    'C': '1', 
    'D': '2', 
    'E': '3', 
    'F': '4', 
    'G': '5', 
    'A': '6', 
    'B': '7'
}

# Mapping fifths value and pitch correction
MAP_CORRECTION = {
    "-7": 0,
    "-6": -4,
    "-5": -1,
    "-4":  -5,
    "-3": -2,
    "-2": -6,
    "-1": -3,
    "0": 0,
    "1": -4,
    "2": -1,
    "3": -5,
    "4": -2,
    "5": -6,
    "6": -3,
    "7": 0,
}

def convert_to_jianpu(note, attributes):
    note_step = note["step"]
    octave = note['octave']
    note_type = note['type']
    duration = note['duration']
    dot_count = note['dot_count']

    fifths = attributes['fifths']
    divisions = attributes['divisions']

    jianpu_note = ""
    note_step_cor = 0

    if note_step in MAP_NOTE:
        if note_step == 'R':
            note_step_cor = 0
            jianpu_note = '0' 
        else:
            temp_step_cor = int(MAP_NOTE[note_step]) + MAP_CORRECTION[str(fifths)]
            note_step_cor = temp_step_cor%7
            
            note_step_cor = note_step_cor if note_step_cor != 0 else 7
            jianpu_note = str(note_step_cor)
            if temp_step_cor <= 0:
                octave -= 1 
    else:
        return ""

    if octave >= 4:
        # 处理高音音符(包括中音音符)
        if octave == 5:
            jianpu_note += "'"
        elif octave == 6:
            jianpu_note += "\""
        elif octave == 7:
            jianpu_note += "`"
        
        # 加入时值标识
        if note_type == 'whole' or duration / divisions == 4:
            jianpu_note += ' - - -'
        elif note_type == 'half' or duration / divisions == 2:
            jianpu_note += ' -'
        elif note_type == 'quarter' or duration / divisions == 1:
            jianpu_note += ''
        elif note_type == 'eighth' or duration / divisions == 0.5:
            jianpu_note += '_'
        elif note_type == "16th" or duration / divisions == 0.25:
            jianpu_note += "="
        elif note_type == "32nd" or duration / divisions == 0.125:
            jianpu_note += "/"
        elif note_type == "64th" or duration / divisions == 1/16:
            jianpu_note += "\\"

    elif octave < 4:
        # TODO: 低音音符处理
        # 下面 copy 了 octave > 4 的处理逻辑，写的时候删去就好。

        if octave == 3:
            if note_type == 'whole' or duration / divisions == 4:
                jianpu_note += 'q - - -'
            elif note_type == 'half' or duration / divisions == 2:
                jianpu_note += 'q -'
            elif note_type == 'quarter' or duration / divisions == 1:
                jianpu_note += 'q' 
            elif note_type == 'eighth' or duration / divisions == 0.5:
                jianpu_note += 'w'
            elif note_type == "16th" or duration / divisions == 0.25:
                jianpu_note += "e"
            elif note_type == "32nd" or duration / divisions == 0.125:
                jianpu_note += "r"
            elif note_type == "64th" or duration / divisions == 1/16:
                jianpu_note += "t"
        elif octave == 2:
            if note_type == 'whole' or duration / divisions == 4:
                jianpu_note += 'a - - -'
            elif note_type == 'half' or duration / divisions == 2:
                jianpu_note += 'a -'
            elif note_type == 'quarter' or duration / divisions == 1:
                jianpu_note += 'a' 
            elif note_type == 'eighth' or duration / divisions == 0.5:
                jianpu_note += 's'
            elif note_type == "16th" or duration / divisions == 0.25:
                jianpu_note += "d"
            elif note_type == "32nd" or duration / divisions == 0.125:
                jianpu_note += "f"
            elif note_type == "64th" or duration / divisions == 1/16:
                jianpu_note += "g"
        elif octave == 1:
            if note_type == 'whole' or duration / divisions == 4:
                jianpu_note += 'z - - -'
            elif note_type == 'half' or duration / divisions == 2:
                jianpu_note += 'z -'
            elif note_type == 'quarter' or duration / divisions == 1:
                jianpu_note += 'z' 
            elif note_type == 'eighth' or duration / divisions == 0.5:
                jianpu_note += 'x'
            elif note_type == "16th" or duration / divisions == 0.25:
                jianpu_note += "c"
            elif note_type == "32nd" or duration / divisions == 0.125:
                jianpu_note += "v"
            elif note_type == "64th" or duration / divisions == 1/16:
                jianpu_note += "g"
    
    # 加入附点   
    if dot_count == 1:
        jianpu_note += '.'
    elif dot_count == 2:
        jianpu_note += '.,'

    return jianpu_note


def parse(file_path) -> str:
    '''返回简谱排版所需要输入的字符串'''
    tree = ET.parse(file_path)
    root = tree.getroot()

    score = ''
    divisions = 1

    fifths = 0
    beam = False

    # 遍历小节(measure)
    for measure in root.findall('.//measure'):
        attributes = measure.find('attributes')
        notes = measure.findall('note') 
        if attributes is not None:
            # 获得乐谱属性
            div = attributes.find('divisions')
            if attributes.find('key/fifths') is not None:
                fifths = (attributes.find('key/fifths').text)
            if div is not None:
                divisions = int(div.text)
        
        # 遍历小节中的音符
        for note in notes:
            rest = note.find('rest')
            pitch = note.find('pitch')
            if rest is not None:
                duration = int(note.find('duration').text)
                note_type = None
                if note.find('type') is not None:
                    note_type = note.find('type').text
                dot_count = len(note.findall('dot'))

                score += convert_to_jianpu(
                            {
                                'step': 'R',
                                'octave': 4,
                                'duration': duration,
                                'type': note_type,
                                'dot_count': dot_count
                            }, 
                            {
                                'fifths': fifths,
                                'divisions': divisions,
                            }
                        ) + ' '

            if pitch is not None:
                step = pitch.find('step').text
                octave = int(pitch.find('octave').text)
                duration = int(note.find('duration').text)
                note_type = note.find('type').text
                dot_count = len(note.findall('dot'))
                beam_be = note.find('beam')
                
                if beam_be is not None:
                    beam = True if beam_be.text == 'begin' else False
                
                spacing = ' '
                if beam:
                    spacing = '' 

                score += convert_to_jianpu(
                            {
                                'step': step,
                                'octave': octave,
                                'duration': duration,
                                'type': note_type,
                                'dot_count': dot_count
                            },
                            {
                                'fifths': fifths,
                                'divisions': divisions,
                            }
                        ) + spacing
        
        # 处理特殊小节线
        barline = measure.find('barline')
        if  barline is not None:
            if barline.find('bar-style').text == "light-light":
                score += "| | "
            elif barline.find('bar-style').text == "light-heavy":
                score += "+"
        else:
            score += "| " 
    return score
    

def create_doc(notes, output_doc):
    doc = Document()
    p = doc.add_paragraph()
    run = p.add_run(notes)
    run.font.name = 'jpfont-nds'
    r = run._element
    rFonts = r.find(qn('w:rPr')).find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = OxmlElement('w:rFonts')
        r.get_or_add('w:rPr').append(rFonts)
    rFonts.set(qn('w:eastAsia'), 'jpfont-nds')
    run.font.size = Pt(12)
    doc.save(output_doc)


if __name__ == "__main__":
    parser = argparse.ArgumentParser()
    parser.add_argument('filename', type=str)
    args = parser.parse_args()

    musicxml_file = args.filename
    output_doc = 'outputs/' + datetime.datetime.now().strftime("%Y_%m_%d_%H_%M_%S") + '.docx'

    notes = parse(musicxml_file)
    create_doc(notes, output_doc)

    print("Score saved to " + output_doc)